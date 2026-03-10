"""Tests for document parser."""

import tempfile
from pathlib import Path

import pytest

from src.documents.parser import (
    extract_text_from_docx,
    extract_text_from_pdf,
    parse_document,
    parse_text_input,
)


def test_parse_text_input():
    """Pasted text is normalized."""
    assert parse_text_input("  hello world  ") == "hello world"
    assert parse_text_input("") == ""
    assert parse_text_input(None) == ""


def test_parse_document_unsupported_format():
    """Unsupported extension raises ValueError."""
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
        f.write(b"hello")
        path = f.name
    try:
        with pytest.raises(ValueError, match="Unsupported document format"):
            parse_document(path)
    finally:
        Path(path).unlink(missing_ok=True)


def test_parse_document_file_not_found():
    """Missing file raises FileNotFoundError for PDF."""
    with pytest.raises(FileNotFoundError, match="not found"):
        parse_document("/nonexistent/file.pdf")


def test_extract_text_from_pdf():
    """PDF with one page returns its text."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        path = f.name
    try:
        # Create minimal PDF with PyMuPDF
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello PDF World")
        doc.save(path)
        doc.close()
        text = extract_text_from_pdf(path)
        assert "Hello PDF World" in text
    finally:
        Path(path).unlink(missing_ok=True)


def test_extract_text_from_docx():
    """DOCX returns paragraph text."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("Hello DOCX World")
        doc.save(path)
        text = extract_text_from_docx(path)
        assert "Hello DOCX World" in text
    finally:
        Path(path).unlink(missing_ok=True)


def test_parse_document_pdf():
    """parse_document dispatches to PDF for .pdf."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        path = f.name
    try:
        import fitz
        doc = fitz.open()
        doc.new_page().insert_text((72, 72), "Parsed PDF")
        doc.save(path)
        doc.close()
        assert "Parsed PDF" in parse_document(path)
    finally:
        Path(path).unlink(missing_ok=True)


def test_parse_document_docx():
    """parse_document dispatches to DOCX for .docx."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("Parsed DOCX")
        doc.save(path)
        assert "Parsed DOCX" in parse_document(path)
    finally:
        Path(path).unlink(missing_ok=True)
